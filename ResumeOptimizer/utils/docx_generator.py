from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Dict, Any, List
import json

class DocxGenerator:
    @staticmethod
    def generate(resume_data: Dict[str, Any], output_path: str):
        doc = Document()

        DocxGenerator._set_default_font(doc)

        if 'personal_info' in resume_data and resume_data['personal_info']:
            DocxGenerator._add_title(doc, "个人简历")

            DocxGenerator._add_section(doc, "基本信息")
            personal_info = resume_data['personal_info']

            if isinstance(personal_info, dict):
                for key, value in personal_info.items():
                    if value and value not in ['null', 'None', '']:
                        DocxGenerator._add_info_line(doc, key, value)

            if 'career_summary' in resume_data and resume_data['career_summary']:
                DocxGenerator._add_section(doc, "职业概述")
                DocxGenerator._add_paragraph(doc, resume_data['career_summary'])

        if 'education' in resume_data and resume_data['education']:
            DocxGenerator._add_section(doc, "教育背景")
            education_list = resume_data['education']

            if isinstance(education_list, list):
                for edu in education_list:
                    if isinstance(edu, dict):
                        school = edu.get('学校名称', '')
                        degree = edu.get('学历学位', '')
                        major = edu.get('专业', '')
                        period = edu.get('起止时间', '')
                        extra = edu.get('附加信息', '')

                        header = f"{school} | {degree} | {major}"
                        if period:
                            header += f" | {period}"
                        DocxGenerator._add_subsection(doc, header)

                        if major:
                            DocxGenerator._add_paragraph(doc, f"专业：{major}", bold=False)
                        if extra:
                            DocxGenerator._add_paragraph(doc, extra)

        if 'experience' in resume_data and resume_data['experience']:
            DocxGenerator._add_section(doc, "工作经历")
            experience_list = resume_data['experience']

            if isinstance(experience_list, list):
                for exp in experience_list:
                    if isinstance(exp, dict):
                        company = exp.get('公司名称', '')
                        position = exp.get('职位名称', '')
                        period = exp.get('在职时间', '')

                        header = f"{company}"
                        if position:
                            header += f" | {position}"
                        if period:
                            header += f" | {period}"
                        DocxGenerator._add_subsection(doc, header)

                        descriptions = exp.get('工作描述', [])
                        if isinstance(descriptions, list):
                            for desc in descriptions:
                                if desc:
                                    DocxGenerator._add_bullet(doc, desc)

                        achievements = exp.get('关键成果', '')
                        if achievements:
                            DocxGenerator._add_paragraph(doc, f"📊 关键成果：{achievements}")

        if 'projects' in resume_data and resume_data['projects']:
            DocxGenerator._add_section(doc, "项目经验")
            project_list = resume_data['projects']

            if isinstance(project_list, list):
                for proj in project_list:
                    if isinstance(proj, dict):
                        name = proj.get('项目名称', '')
                        period = proj.get('项目时间', '')

                        header = f"{name}"
                        if period:
                            header += f" | {period}"
                        DocxGenerator._add_subsection(doc, header)

                        description = proj.get('项目描述', '')
                        if description:
                            DocxGenerator._add_paragraph(doc, description)

                        responsibility = proj.get('个人职责', '')
                        if responsibility:
                            DocxGenerator._add_paragraph(doc, f"职责：{responsibility}")

                        tech_stack = proj.get('技术栈', '')
                        if tech_stack:
                            if isinstance(tech_stack, list):
                                tech_stack = '、'.join(tech_stack)
                            DocxGenerator._add_paragraph(doc, f"技术栈：{tech_stack}")

                        result = proj.get('项目成果', '')
                        if result:
                            DocxGenerator._add_paragraph(doc, f"📊 项目成果：{result}")

        if 'skills' in resume_data and resume_data['skills']:
            DocxGenerator._add_section(doc, "技能证书")
            skills_data = resume_data['skills']

            if isinstance(skills_data, dict):
                professional_skills = skills_data.get('专业技能', [])
                if professional_skills:
                    if isinstance(professional_skills, list):
                        skill_texts = []
                        for skill in professional_skills:
                            if isinstance(skill, dict):
                                name = skill.get('技能名称', '')
                                level = skill.get('熟练度', '')
                                if name:
                                    skill_texts.append(f"{name}（{level}）" if level else name)
                            elif isinstance(skill, str):
                                skill_texts.append(skill)
                        if skill_texts:
                            DocxGenerator._add_paragraph(doc, "专业技能：")
                            for skill_text in skill_texts:
                                DocxGenerator._add_bullet(doc, skill_text)

                language_skills = skills_data.get('语言能力', [])
                if language_skills:
                    if isinstance(language_skills, list):
                        for lang in language_skills:
                            if isinstance(lang, dict):
                                language = lang.get('语言', '')
                                level = lang.get('水平', '')
                                if language:
                                    DocxGenerator._add_bullet(doc, f"{language}：{level}" if level else language)
                            elif isinstance(lang, str):
                                DocxGenerator._add_bullet(doc, lang)

                certificates = skills_data.get('证书认证', [])
                if certificates:
                    if isinstance(certificates, list):
                        DocxGenerator._add_paragraph(doc, "证书认证：")
                        for cert in certificates:
                            if isinstance(cert, str):
                                DocxGenerator._add_bullet(doc, cert)

        if 'additional_value' in resume_data and resume_data['additional_value']:
            DocxGenerator._add_section(doc, "附加价值")
            add_value = resume_data['additional_value']
            if isinstance(add_value, dict):
                for key, value in add_value.items():
                    if value:
                        DocxGenerator._add_bullet(doc, f"{key}：{value}")

        industry_analysis = resume_data.get('industry_analysis', '')
        if industry_analysis:
            DocxGenerator._add_section(doc, "行业分析")
            DocxGenerator._add_paragraph(doc, industry_analysis)

        competitiveness = resume_data.get('competitiveness', {})
        if competitiveness and isinstance(competitiveness, dict):
            DocxGenerator._add_section(doc, "竞争力评估")
            score = competitiveness.get('overall_score', '')
            if score:
                DocxGenerator._add_paragraph(doc, f"综合评分：{score}/100")
            analysis = competitiveness.get('analysis', '')
            if analysis:
                DocxGenerator._add_paragraph(doc, analysis)

        career_suggestions = resume_data.get('career_suggestions', {})
        if career_suggestions and isinstance(career_suggestions, dict):
            DocxGenerator._add_section(doc, "职业发展建议")

            directions = career_suggestions.get('就业方向', [])
            if directions:
                DocxGenerator._add_subsection(doc, "建议就业方向")
                for d in directions:
                    if isinstance(d, str) and d:
                        DocxGenerator._add_bullet(doc, d)

            skills_to_add = career_suggestions.get('补充技能', [])
            if skills_to_add:
                DocxGenerator._add_subsection(doc, "建议补充技能")
                for s in skills_to_add:
                    if isinstance(s, str) and s:
                        DocxGenerator._add_bullet(doc, s)

            actions = career_suggestions.get('短期行动', [])
            if actions:
                DocxGenerator._add_subsection(doc, "短期行动建议")
                for a in actions:
                    if isinstance(a, str) and a:
                        DocxGenerator._add_bullet(doc, a)

        humanized_message = resume_data.get('humanized_message', '')
        if humanized_message:
            DocxGenerator._add_section(doc, "给你的话")
            DocxGenerator._add_paragraph(doc, humanized_message)

        doc.save(output_path)

    @staticmethod
    def _set_default_font(doc: Document):
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    @staticmethod
    def _add_title(doc: Document, text: str):
        title = doc.add_heading(text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)

    @staticmethod
    def _add_section(doc: Document, text: str):
        heading = doc.add_heading(text, level=1)
        for run in heading.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 102, 153)

    @staticmethod
    def _add_subsection(doc: Document, text: str):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(51, 51, 51)
        return para

    @staticmethod
    def _add_info_line(doc: Document, label: str, value: str):
        para = doc.add_paragraph()
        label_run = para.add_run(f"{label}：")
        label_run.font.name = '微软雅黑'
        label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        label_run.bold = True
        label_run.font.size = Pt(10.5)

        value_run = para.add_run(value)
        value_run.font.name = '微软雅黑'
        value_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        value_run.font.size = Pt(10.5)

    @staticmethod
    def _add_paragraph(doc: Document, text: str, bold: bool = False):
        para = doc.add_paragraph(text)
        for run in para.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.bold = bold
            run.font.size = Pt(10.5)
        return para

    @staticmethod
    def _add_bullet(doc: Document, text: str):
        para = doc.add_paragraph(text, style='List Bullet')
        for run in para.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(10.5)
        return para
