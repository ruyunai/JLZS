import pdfplumber
from docx import Document
from typing import Dict, Any, Optional
import re

class FileParser:
    @staticmethod
    def parse_resume(file_path: str) -> Dict[str, Any]:
        file_path_lower = file_path.lower()
        if file_path_lower.endswith('.pdf'):
            return FileParser._parse_pdf(file_path)
        elif file_path_lower.endswith('.docx'):
            return FileParser._parse_docx(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_path}")

    @staticmethod
    def _parse_pdf(file_path: str) -> Dict[str, Any]:
        text_content = []
        tables_content = []

        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)

                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            table_text = '\n'.join([' | '.join([str(cell) if cell else '' for cell in row]) for row in table])
                            tables_content.append(table_text)
        except Exception as e:
            raise Exception(f"PDF解析失败: {str(e)}")

        full_text = '\n\n'.join(text_content)
        return {
            'text': full_text,
            'tables': tables_content,
            'raw_content': text_content
        }

    @staticmethod
    def _parse_docx(file_path: str) -> Dict[str, Any]:
        try:
            doc = Document(file_path)
            full_text = '\n'.join([para.text for para in doc.paragraphs])
            tables_content = []

            for table in doc.tables:
                table_text = '\n'.join([' | '.join([cell.text for cell in row.cells]) for row in table.rows])
                tables_content.append(table_text)

            return {
                'text': full_text,
                'tables': tables_content,
                'raw_content': [para.text for para in doc.paragraphs]
            }
        except Exception as e:
            raise Exception(f"Word文档解析失败: {str(e)}")

    @staticmethod
    def extract_structured_info(text: str) -> Dict[str, Any]:
        return {
            'full_text': text,
            'has_contact': bool(re.search(r'1[3-9]\d{9}|[\w\.-]+@[\w\.-]+', text)),
            'has_education': bool(re.search(r'大学|学院|本科|硕士|博士|大专', text)),
            'has_experience': bool(re.search(r'工作|实习|项目|经验', text)),
            'word_count': len(text)
        }
